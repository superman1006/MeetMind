"""文档加载器：把 JSON / Markdown / PDF / Word / 纯文本 文件统一解析为 RAG 文档列表。

每种文件格式对应一个独立函数，输入是 `Path`，输出是 `list[dict]`；
返回的 dict 至少包含 `content` 字段，可选 `type`、`source`。
顶层分发函数 `load_file()` 根据文件后缀决定调用哪个 loader。
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Callable

from meetmind.utils.logger import get_logger

logger = get_logger(__name__)


# 加载 JSON
def load_json(path: Path) -> list[dict]:
    """加载 JSON 文件。

    要求文件根是 `list[dict]`，每个 dict 至少有 `content` 字段，
    其余字段（type / date / source 等）会作为 metadata 一起保留。
    """
    raw = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(raw, list):
        logger.warning(f"JSON {path} 顶层不是 list，已跳过")
        return []
    docs: list[dict] = []

    # 拿到 list 后逐条检查，丢弃不符合要求的条目（非 dict 或缺 content）
    for i, item in enumerate(raw):
        if not isinstance(item, dict) or "content" not in item:
            continue
        item.setdefault("type", "json")
        item.setdefault("source", path.name)
        docs.append(item)
    return docs


# 加载 Markdown
def load_markdown(path: Path) -> list[dict]:
    """加载 Markdown 文件，按二级及以下标题或空行切块。

    简单策略：以 `##` 标题为分块边界；若没有标题，则按空行段落分块。
    每一块作为一条 RAG 文档。
    """
    text = path.read_text(encoding="utf-8")
    blocks: list[str] = []
    current: list[str] = []
    has_heading = False
    # 按行遍历，遇到 `##` 标题就切块；否则继续累积到 current
    for line in text.splitlines():
        if line.startswith("## "):
            has_heading = True
            if current:
                blocks.append("\n".join(current).strip())
            current = [line]
        else:
            current.append(line)
    if current:
        blocks.append("\n".join(current).strip())

    # 没有任何 `##` 标题 → 退化为按空行分段
    if not has_heading:
        blocks = []
        for b in text.split("\n\n"):
            blocks.append(b.strip())

    result = []
    for b in blocks:
        if b.strip():
            result.append({"content": b, "type": "markdown", "source": path.name})
    return result


# 加载 PDF
def load_pdf(path: Path) -> list[dict]:
    """加载 PDF 文件，每页作为一条 RAG 文档。

    使用 `pypdf` 提取文本。扫描版 PDF（图像）将得到空字符串并被丢弃。
    """
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    docs: list[dict] = []
    # 把每页作为一个 doc，丢弃空页
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text().strip()
        if not text:
            continue
        docs.append(
            {
                "content": text,
                "type": "pdf",
                "source": f"{path.name}#page{i}",
            }
        )
    return docs


# 加载 Word
def load_docx(path: Path) -> list[dict]:
    """加载 Word (.docx) 文件，按段落作为一条 RAG 文档。

    跳过空段落；表格内容暂不处理（如有需要可扩展遍历 `doc.tables`）。
    """
    from docx import Document

    document = Document(str(path))
    docs: list[dict] = []

    # 把每个段落作为一个 doc，丢弃空段落
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        docs.append({"content": text, "type": "docx", "source": path.name})
    return docs


# 加载 纯文本
def load_text(path: Path) -> list[dict]:
    """加载 .txt 等纯文本文件，按空行分段。"""
    text : str= path.read_text(encoding="utf-8")
    raw_chunks : list[str] = text.split("\n\n")

    # 把切块后的每块文本去掉首尾空白，丢弃空块
    chunks = []
    for c in raw_chunks:
        stripped = c.strip()
        if stripped:
            chunks.append(stripped)

    result : list[dict] = []
    for c in chunks:
        result.append({"content": c, "type": "text", "source": path.name})
    return result


# 分发器
LOADERS: dict[str, Callable[[Path], list[dict]]] = {
    ".json": load_json,
    ".md": load_markdown,
    ".markdown": load_markdown,
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".txt": load_text,
}


def load_file(path: Path) -> list[dict]:
    """根据文件后缀分发到对应 loader；返回类型为 list[dict]，每个 dict 至少有 `content` 字段。"""
    loader : Callable = LOADERS.get(path.suffix.lower())
    if loader is None:
        logger.warning(f"不支持的文件类型，跳过：{path.name}")
        return []

    try:
        return loader(path)
    except Exception as exc:
        logger.error(f"解析 {path} 失败：{exc}")
        return []
